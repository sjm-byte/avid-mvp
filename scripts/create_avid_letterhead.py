# -*- coding: utf-8 -*-
"""Generate Avid A4 letterhead matching the modern navy/cyan corporate template.

Source of truth: letterhe*ad-13.jpg from
  modern-business-corporate-letterhead-template_9
PSD is not opened; decorative geometry is redrawn from measured JPG colors
and proportions. All written content is editable Word text.

Outputs (Desktop):
  - Letterhead-A4.docx
  - Letterhead-A4-preview.png
"""

from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from PIL import Image, ImageDraw

# ---------------------------------------------------------------------------
# Colors measured from letterhe*ad-13.jpg
# ---------------------------------------------------------------------------

NAVY = (18, 55, 74)  # #12374A
CYAN = (3, 194, 239)  # #03C2EF
GREY = (120, 128, 136)
SOFT = (90, 100, 110)
WHITE = (255, 255, 255)

NAVY_RGB = RGBColor(*NAVY)
CYAN_RGB = RGBColor(*CYAN)
GREY_RGB = RGBColor(*GREY)
SOFT_RGB = RGBColor(*SOFT)

REPO = Path(r"c:\Users\Almahdi Laptop\OneDrive\AVID\Avid-MVP")
ASSETS = REPO / "scripts" / "_letterhead_assets"
LOGO_MARK = REPO / "public" / "assets" / "brand" / "minilogo.png"
LOGO_HEADER = REPO / "public" / "assets" / "brand" / "logoheader.png"

# Editable defaults (Persian + Latin, matching template slots)
COMPANY_FA = "آوید"
COMPANY_EN = "AVID"
TAGLINE = "مدیریت مشارکت در طرح‌های سرمایه‌گذاری"
SENDER_NAME = "نام فرستنده"
SENDER_TITLE = "سمت سازمانی"
PHONE = "۰۲۱-۱۲۳۴۵۶۷۸"
MOBILE = "۰۹۱۲۱۲۳۴۵۶۷"
WEB = "avidinvest.ir"
EMAIL = "info@avidinvest.ir"
ADDR_LINE1 = "تهران، خیابان نمونه، پلاک ۱۲"
ADDR_LINE2 = "ایران"
DATE_LABEL = "تاریخ :"
DATE_VALUE = "…………"

# A4 content width used for deco images (matches Word margins 1.75 cm each side)
CONTENT_WIDTH_CM = 17.5
HEADER_STUB_CM = 1.25
HEADER_BRAND_CM = 4.05
HEADER_BARS_CM = 12.2
HEADER_DECO_H_CM = 1.45
FOOTER_DECO_H_CM = 2.05
DPI = 300


def resolve_desktop() -> Path:
    primary = Path(r"C:\Users\Almahdi Laptop\Desktop")
    if primary.is_dir():
        return primary
    desktop = Path.home() / "Desktop"
    try:
        import subprocess

        out = subprocess.check_output(
            [
                "powershell",
                "-NoProfile",
                "-Command",
                "[Environment]::GetFolderPath('Desktop')",
            ],
            text=True,
            encoding="utf-8",
            errors="replace",
        ).strip()
        if out:
            desktop = Path(out)
    except Exception:
        pass
    return desktop


def cm_to_px(cm: float, dpi: int = DPI) -> int:
    return max(1, int(round(cm / 2.54 * dpi)))


# ---------------------------------------------------------------------------
# Geometry helpers (rotated square / diamond)
# ---------------------------------------------------------------------------

def diamond_points(cx: float, cy: float, half_diag: float) -> list[tuple[float, float]]:
    """Axis-aligned bbox diamond = square rotated 45°; half_diag is half of bbox width."""
    return [
        (cx, cy - half_diag),
        (cx + half_diag, cy),
        (cx, cy + half_diag),
        (cx - half_diag, cy),
    ]


def punch_diamond(
    draw: ImageDraw.ImageDraw,
    cx: float,
    cy: float,
    half_diag: float,
    border: float,
    fill_rgb: tuple[int, int, int],
):
    """Clear a white halo, then paint cyan diamond (breaks navy bands)."""
    outer = diamond_points(cx, cy, half_diag + border)
    draw.polygon(outer, fill=WHITE)
    inner = diamond_points(cx, cy, half_diag)
    draw.polygon(inner, fill=fill_rgb)
    # Thin white rim for separation (draw slightly larger white then cyan again)
    rim = max(2.0, border * 0.35)
    draw.polygon(diamond_points(cx, cy, half_diag + rim), outline=WHITE)
    draw.polygon(inner, fill=fill_rgb)


def draw_twin_bands(
    draw: ImageDraw.ImageDraw,
    x0: float,
    x1: float,
    y_top: float,
    band_h: float,
    gap: float,
):
    """Two parallel navy rectangles from x0..x1."""
    draw.rectangle([x0, y_top, x1, y_top + band_h], fill=NAVY)
    y2 = y_top + band_h + gap
    draw.rectangle([x0, y2, x1, y2 + band_h], fill=NAVY)


# ---------------------------------------------------------------------------
# Decorative PNG builders
# ---------------------------------------------------------------------------

def build_header_stub_png(path: Path) -> Path:
    """Left stub of twin navy bands (logo gap sits to its right)."""
    w = cm_to_px(HEADER_STUB_CM)
    h = cm_to_px(HEADER_DECO_H_CM)
    img = Image.new("RGBA", (w, h), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)
    # JPG: band_h : gap : band_h ≈ 80 : 40 : 80; diamond overhang above top band
    pad_top = int(h * 0.16)
    band_h = int(h * 0.28)
    gap = int(h * 0.14)
    draw_twin_bands(draw, 0, w - 1, pad_top, band_h, gap)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, "PNG")
    return path


def build_header_bars_png(path: Path) -> Path:
    """Right header twin bands + cyan diamond on the right (no text)."""
    w = cm_to_px(HEADER_BARS_CM)
    h = cm_to_px(HEADER_DECO_H_CM)
    img = Image.new("RGBA", (w, h), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)

    pad_top = int(h * 0.16)
    band_h = int(h * 0.28)
    gap = int(h * 0.14)
    bands_bottom = pad_top + band_h + gap + band_h
    cy = pad_top + (bands_bottom - pad_top) / 2

    # Diamond toward the right of this strip (JPG ~62% of right piece ≈ page-right)
    cx = int(w * 0.70)
    half = int(h * 0.38)
    border = max(4, int(h * 0.075))

    left_end = cx - half - border - 1
    right_start = cx + half + border + 1
    if left_end > 2:
        draw_twin_bands(draw, 0, left_end, pad_top, band_h, gap)
    if right_start < w - 2:
        draw_twin_bands(draw, right_start, w - 1, pad_top, band_h, gap)

    punch_diamond(draw, cx, cy, half, border, CYAN)

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, "PNG")
    return path


def build_footer_bars_png(path: Path) -> Path:
    """Full-width footer twin bands + centered cyan diamond (no text)."""
    w = cm_to_px(CONTENT_WIDTH_CM)
    h = cm_to_px(FOOTER_DECO_H_CM)
    img = Image.new("RGBA", (w, h), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)

    pad_top = int(h * 0.08)
    band_h = int(h * 0.32)
    gap = int(h * 0.16)
    bands_bottom = pad_top + band_h + gap + band_h
    cy = pad_top + (bands_bottom - pad_top) / 2
    cx = int(w * 0.48)
    half = int(h * 0.44)
    border = max(5, int(h * 0.085))

    left_end = cx - half - border - 1
    right_start = cx + half + border + 1
    margin = int(w * 0.008)
    if left_end > margin:
        draw_twin_bands(draw, margin, left_end, pad_top, band_h, gap)
    if right_start < w - margin:
        draw_twin_bands(draw, right_start, w - margin - 1, pad_top, band_h, gap)

    punch_diamond(draw, cx, cy, half, border, CYAN)

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, "PNG")
    return path


def build_circular_icon(path: Path, kind: str, size: int = 96) -> Path:
    img = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)
    draw.ellipse([1, 1, size - 2, size - 2], fill=NAVY)
    pad = size * 0.28
    box = (pad, pad, size - pad, size - pad)
    x0, y0, x1, y1 = box
    cx, cy = (x0 + x1) / 2, (y0 + y1) / 2
    s = x1 - x0
    stroke = max(2, int(s / 9))
    c = WHITE

    if kind == "phone":
        # Simple handset silhouette
        draw.rounded_rectangle(
            [cx - s * 0.22, cy - s * 0.38, cx + s * 0.22, cy + s * 0.38],
            radius=s * 0.12,
            outline=c,
            width=stroke,
        )
        draw.line(
            [cx - s * 0.1, cy - s * 0.28, cx + s * 0.1, cy - s * 0.28],
            fill=c,
            width=stroke,
        )
    elif kind == "web":
        r = s * 0.38
        draw.ellipse([cx - r, cy - r, cx + r, cy + r], outline=c, width=stroke)
        draw.ellipse(
            [cx - r * 0.42, cy - r, cx + r * 0.42, cy + r],
            outline=c,
            width=max(1, stroke - 1),
        )
        draw.line([cx - r, cy, cx + r, cy], fill=c, width=max(1, stroke - 1))
        draw.line(
            [cx - r * 0.85, cy - r * 0.4, cx + r * 0.85, cy - r * 0.4],
            fill=c,
            width=1,
        )
        draw.line(
            [cx - r * 0.85, cy + r * 0.4, cx + r * 0.85, cy + r * 0.4],
            fill=c,
            width=1,
        )
    elif kind == "location":
        r = s * 0.28
        draw.ellipse(
            [cx - r, cy - r - s * 0.1, cx + r, cy + r - s * 0.1],
            outline=c,
            width=stroke,
        )
        draw.polygon(
            [
                (cx - r * 0.9, cy + r * 0.2 - s * 0.05),
                (cx + r * 0.9, cy + r * 0.2 - s * 0.05),
                (cx, y1 - 1),
            ],
            fill=c,
        )
        draw.ellipse(
            [
                cx - r * 0.35,
                cy - r * 0.35 - s * 0.1,
                cx + r * 0.35,
                cy + r * 0.35 - s * 0.1,
            ],
            fill=NAVY,
        )

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, "PNG")
    return path


def build_avid_logo_png(path: Path, size: int = 160) -> Path:
    """Cyan geometric mark for the template logo slot (Avid branding)."""
    img = Image.new("RGBA", (size, size), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)
    m = int(size * 0.06)
    t = max(3, size // 12)
    # four L-corners (template-like broken square)
    draw.rectangle([m, m, m + t, size // 2 - m], fill=CYAN)
    draw.rectangle([m, m, size // 2 - m, m + t], fill=CYAN)
    draw.rectangle([size - m - t, m, size - m, size // 2 - m], fill=CYAN)
    draw.rectangle([size // 2 + m, m, size - m, m + t], fill=CYAN)
    draw.rectangle([m, size // 2 + m, m + t, size - m], fill=CYAN)
    draw.rectangle([m, size - m - t, size // 2 - m, size - m], fill=CYAN)
    draw.rectangle([size - m - t, size // 2 + m, size - m, size - m], fill=CYAN)
    draw.rectangle([size // 2 + m, size - m - t, size - m, size - m], fill=CYAN)

    src = LOGO_HEADER if LOGO_HEADER.is_file() else LOGO_MARK
    if src.is_file():
        logo = Image.open(src).convert("RGBA")
        # Downscale first for speed, then chroma-key near-black
        logo.thumbnail((size, size), Image.Resampling.LANCZOS)
        px = logo.load()
        lw, lh = logo.size
        for y in range(lh):
            for x in range(lw):
                r, g, b, a = px[x, y]
                lum = int(0.299 * r + 0.587 * g + 0.114 * b)
                if a < 16 or lum < 25:
                    px[x, y] = (0, 0, 0, 0)
                else:
                    px[x, y] = (*NAVY, min(255, int(a * (0.6 + lum / 255 * 0.4))))
        bbox = logo.getbbox()
        if bbox:
            logo = logo.crop(bbox)
        inner = int(size * 0.52)
        logo = logo.resize((inner, inner), Image.Resampling.LANCZOS)
        tile = Image.new("RGBA", (inner + 8, inner + 8), (0, 0, 0, 0))
        td = ImageDraw.Draw(tile)
        td.rounded_rectangle(
            [0, 0, inner + 7, inner + 7], radius=inner // 6, fill=CYAN
        )
        ox = (size - tile.width) // 2
        oy = (size - tile.height) // 2
        img.paste(tile, (ox, oy), tile)
        img.paste(logo, (ox + 4, oy + 4), logo)
    else:
        draw.rounded_rectangle(
            [size * 0.28, size * 0.28, size * 0.72, size * 0.72],
            radius=size * 0.08,
            fill=CYAN,
        )

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, "PNG")
    return path


def generate_decorative_assets() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    return {
        "header_stub": build_header_stub_png(ASSETS / "deco_header_stub.png"),
        "header_bars": build_header_bars_png(ASSETS / "deco_header_bars.png"),
        "footer_bars": build_footer_bars_png(ASSETS / "deco_footer_bars.png"),
        "logo": build_avid_logo_png(ASSETS / "deco_logo_avid.png"),
        "icon_phone": build_circular_icon(ASSETS / "deco_icon_phone_circle.png", "phone"),
        "icon_web": build_circular_icon(ASSETS / "deco_icon_web_circle.png", "web"),
        "icon_location": build_circular_icon(
            ASSETS / "deco_icon_location_circle.png", "location"
        ),
    }


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def set_run_font(
    run,
    *,
    size_pt: float,
    bold: bool = False,
    color: RGBColor = NAVY_RGB,
    rtl: bool = False,
    font_name: str = "Calibri",
):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    if rtl:
        rtl_el = OxmlElement("w:rtl")
        rtl_el.set(qn("w:val"), "1")
        rPr.append(rtl_el)


def set_paragraph_format(paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, rtl: bool = False):
    paragraph.alignment = align
    pPr = paragraph._p.get_or_add_pPr()
    if rtl:
        bidi = OxmlElement("w:bidi")
        bidi.set(qn("w:val"), "1")
        pPr.append(bidi)
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def clear_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def set_table_fixed_layout(table, widths_cm: list[float]):
    table.autofit = False
    table.allow_autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    existing = tblPr.find(qn("w:tblLayout"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(layout)
    total = sum(widths_cm)
    table.width = Cm(total)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[idx])


def set_cell_margins(cell, top=40, bottom=40, left=40, right=40):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_vertical_align(cell, val: str = "center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), val)
    existing = tcPr.find(qn("w:vAlign"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(vAlign)


def remove_cell_paragraphs(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag == qn("w:p"):
            tc.remove(child)


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def _remove_leading_empty_paragraphs(container):
    """Drop blank leading paragraphs so tables sit flush at the top."""
    element = container._element
    for child in list(element):
        if child.tag != qn("w:p"):
            break
        texts = [t.text for t in child.iter(qn("w:t")) if t.text]
        drawings = list(child.iter(qn("w:drawing")))
        if texts or drawings:
            break
        element.remove(child)


def build_header(section, assets: dict[str, Path]):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    set_paragraph_format(hp, WD_ALIGN_PARAGRAPH.LEFT)

    # Row: stub | logo+name | bars+diamond
    widths = [HEADER_STUB_CM, HEADER_BRAND_CM, HEADER_BARS_CM]
    table = header.add_table(rows=1, cols=3, width=Cm(sum(widths)))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(table)
    set_table_fixed_layout(table, widths)
    _remove_leading_empty_paragraphs(header)

    stub_cell = table.cell(0, 0)
    brand_cell = table.cell(0, 1)
    bars_cell = table.cell(0, 2)
    set_cell_margins(stub_cell, 0, 0, 0, 0)
    set_cell_margins(brand_cell, 30, 10, 40, 20)
    set_cell_margins(bars_cell, 0, 0, 0, 0)
    set_cell_vertical_align(stub_cell, "center")
    set_cell_vertical_align(brand_cell, "center")
    set_cell_vertical_align(bars_cell, "center")

    sp = stub_cell.paragraphs[0]
    sp.clear()
    set_paragraph_format(sp)
    sp.add_run().add_picture(
        str(assets["header_stub"]),
        width=Cm(HEADER_STUB_CM),
        height=Cm(HEADER_DECO_H_CM),
    )

    bp = brand_cell.paragraphs[0]
    bp.clear()
    set_paragraph_format(bp, WD_ALIGN_PARAGRAPH.LEFT)
    logo_run = bp.add_run()
    logo_run.add_picture(str(assets["logo"]), width=Cm(0.9), height=Cm(0.9))
    gap = bp.add_run("  ")
    set_run_font(gap, size_pt=14, color=NAVY_RGB)
    # Template pattern: first word navy, second cyan → آوید + AVID
    r1 = bp.add_run(COMPANY_FA)
    set_run_font(r1, size_pt=15, bold=True, color=NAVY_RGB, rtl=True, font_name="Tahoma")
    r2 = bp.add_run(" ")
    set_run_font(r2, size_pt=13, color=CYAN_RGB)
    r3 = bp.add_run(COMPANY_EN)
    set_run_font(r3, size_pt=13, bold=True, color=CYAN_RGB, font_name="Calibri")

    tag = brand_cell.add_paragraph()
    set_paragraph_format(tag, WD_ALIGN_PARAGRAPH.LEFT, rtl=True)
    tag.paragraph_format.space_before = Pt(1)
    tr = tag.add_run(TAGLINE)
    set_run_font(tr, size_pt=7.5, color=GREY_RGB, rtl=True, font_name="Tahoma")

    bar_p = bars_cell.paragraphs[0]
    bar_p.clear()
    set_paragraph_format(bar_p, WD_ALIGN_PARAGRAPH.RIGHT)
    bar_p.add_run().add_picture(
        str(assets["header_bars"]),
        width=Cm(HEADER_BARS_CM),
        height=Cm(HEADER_DECO_H_CM),
    )

    # Sender + date row
    meta = header.add_table(rows=1, cols=2, width=Cm(CONTENT_WIDTH_CM))
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(meta)
    set_table_fixed_layout(meta, [11.5, 6.0])
    left = meta.cell(0, 0)
    right = meta.cell(0, 1)
    set_cell_margins(left, 120, 20, 40, 40)
    set_cell_margins(right, 120, 20, 40, 40)
    set_cell_vertical_align(right, "bottom")

    p = left.paragraphs[0]
    p.clear()
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT, rtl=True)
    rn = p.add_run(SENDER_NAME)
    set_run_font(rn, size_pt=12, bold=True, color=CYAN_RGB, rtl=True, font_name="Tahoma")

    p2 = left.add_paragraph()
    set_paragraph_format(p2, WD_ALIGN_PARAGRAPH.LEFT, rtl=True)
    p2.paragraph_format.space_before = Pt(2)
    rt = p2.add_run(SENDER_TITLE)
    set_run_font(rt, size_pt=9, color=GREY_RGB, rtl=True, font_name="Tahoma")

    for label, value, is_fa in (
        ("PHONE : ", PHONE, True),
        ("WEB : ", WEB, False),
        ("ADDR : ", ADDR_LINE1, True),
    ):
        pr = left.add_paragraph()
        set_paragraph_format(pr, WD_ALIGN_PARAGRAPH.LEFT, rtl=is_fa)
        pr.paragraph_format.space_before = Pt(3)
        rl = pr.add_run(label)
        set_run_font(rl, size_pt=9, bold=True, color=NAVY_RGB, font_name="Calibri")
        rv = pr.add_run(value)
        set_run_font(
            rv,
            size_pt=9,
            color=SOFT_RGB,
            rtl=is_fa,
            font_name="Tahoma" if is_fa else "Calibri",
        )

    dp = right.paragraphs[0]
    dp.clear()
    set_paragraph_format(dp, WD_ALIGN_PARAGRAPH.RIGHT, rtl=True)
    d1 = dp.add_run(DATE_LABEL + " ")
    set_run_font(d1, size_pt=10, bold=True, color=NAVY_RGB, rtl=True, font_name="Tahoma")
    d2 = dp.add_run(DATE_VALUE)
    set_run_font(d2, size_pt=10, color=SOFT_RGB, rtl=True, font_name="Tahoma")


def build_footer(section, assets: dict[str, Path]):
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    set_paragraph_format(fp)

    # Signatory | contacts
    top = footer.add_table(rows=1, cols=2, width=Cm(CONTENT_WIDTH_CM))
    top.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(top)
    set_table_fixed_layout(top, [9.0, 8.5])
    _remove_leading_empty_paragraphs(footer)
    sig = top.cell(0, 0)
    contact = top.cell(0, 1)
    set_cell_margins(sig, 20, 60, 40, 40)
    set_cell_margins(contact, 20, 40, 40, 20)

    sp = sig.paragraphs[0]
    sp.clear()
    set_paragraph_format(sp, WD_ALIGN_PARAGRAPH.LEFT, rtl=True)
    sn = sp.add_run(SENDER_NAME)
    set_run_font(sn, size_pt=11, bold=True, color=CYAN_RGB, rtl=True, font_name="Tahoma")
    st = sig.add_paragraph()
    set_paragraph_format(st, WD_ALIGN_PARAGRAPH.LEFT, rtl=True)
    st.paragraph_format.space_before = Pt(2)
    st.add_run()
    set_run_font(
        st.add_run(SENDER_TITLE),
        size_pt=9,
        color=GREY_RGB,
        rtl=True,
        font_name="Tahoma",
    )
    # Signature placeholder line (editable)
    ss = sig.add_paragraph()
    set_paragraph_format(ss, WD_ALIGN_PARAGRAPH.LEFT)
    ss.paragraph_format.space_before = Pt(8)
    sr = ss.add_run("__________________")
    set_run_font(sr, size_pt=11, color=NAVY_RGB, font_name="Calibri")
    note = sig.add_paragraph()
    set_paragraph_format(note, WD_ALIGN_PARAGRAPH.LEFT, rtl=True)
    set_run_font(
        note.add_run("امضا"),
        size_pt=8,
        color=GREY_RGB,
        rtl=True,
        font_name="Tahoma",
    )

    rows = [
        (assets["icon_phone"], [PHONE, MOBILE], True),
        (assets["icon_web"], [EMAIL, WEB], False),
        (assets["icon_location"], [ADDR_LINE1, ADDR_LINE2], True),
    ]
    first = True
    for icon, lines, is_fa in rows:
        if first:
            p = contact.paragraphs[0]
            p.clear()
            first = False
        else:
            p = contact.add_paragraph()
        set_paragraph_format(p, WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(icon), width=Cm(0.42), height=Cm(0.42))
        gap = p.add_run("  ")
        set_run_font(gap, size_pt=8, color=NAVY_RGB)
        text = "   |   ".join(lines)
        set_run_font(
            p.add_run(text),
            size_pt=8,
            color=SOFT_RGB,
            rtl=is_fa,
            font_name="Tahoma" if is_fa else "Calibri",
        )

    # Decorative footer strip
    deco_p = footer.add_paragraph()
    set_paragraph_format(deco_p, WD_ALIGN_PARAGRAPH.CENTER)
    deco_p.paragraph_format.space_before = Pt(6)
    deco_p.add_run().add_picture(
        str(assets["footer_bars"]), width=Cm(CONTENT_WIDTH_CM), height=Cm(FOOTER_DECO_H_CM)
    )


def build_docx(assets: dict[str, Path], out_docx: Path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(5.2)
    section.bottom_margin = Cm(4.4)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.45)

    build_header(section, assets)
    build_footer(section, assets)

    # Minimal body — user writes the letter
    body = doc.add_paragraph()
    set_paragraph_format(body, WD_ALIGN_PARAGRAPH.RIGHT, rtl=True)
    body.paragraph_format.space_before = Pt(12)
    set_run_font(
        body.add_run("با سلام و احترام،"),
        size_pt=12,
        color=NAVY_RGB,
        rtl=True,
        font_name="Tahoma",
    )
    for _ in range(12):
        p = doc.add_paragraph()
        set_paragraph_format(p, WD_ALIGN_PARAGRAPH.RIGHT, rtl=True)
        p.paragraph_format.space_after = Pt(12)

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


# ---------------------------------------------------------------------------
# Preview PNG
# ---------------------------------------------------------------------------

def build_preview_png(assets: dict[str, Path], out_path: Path):
    import arabic_reshaper
    from bidi.algorithm import get_display
    from PIL import ImageFont

    def reshape(t: str) -> str:
        return get_display(arabic_reshaper.reshape(t))

    def font(path: Path, size: int):
        try:
            return ImageFont.truetype(str(path), size)
        except OSError:
            return ImageFont.load_default()

    W, H = 1240, 1754  # ~150 DPI A4
    img = Image.new("RGB", (W, H), WHITE)
    draw = ImageDraw.Draw(img)
    margin = int(W * 0.08)

    # Header deco: stub + brand gap + bars
    stub = Image.open(assets["header_stub"]).convert("RGBA")
    bars = Image.open(assets["header_bars"]).convert("RGBA")
    logo = Image.open(assets["logo"]).convert("RGBA")

    deco_h = int(H * 0.052)
    stub_w = int(W * 0.065)
    brand_w = int(W * 0.21)
    bars_w = W - 2 * margin - stub_w - brand_w
    y0 = int(H * 0.028)

    stub_r = stub.resize((stub_w, deco_h), Image.Resampling.LANCZOS)
    bars_r = bars.resize((bars_w, deco_h), Image.Resampling.LANCZOS)
    img.paste(stub_r, (margin, y0), stub_r)

    lh = int(deco_h * 0.68)
    logo_r = logo.resize((lh, lh), Image.Resampling.LANCZOS)
    lx = margin + stub_w + 6
    ly = y0 + (deco_h - lh) // 2
    img.paste(logo_r, (lx, ly), logo_r)

    fa_b = font(Path(r"C:\Windows\Fonts\tahomabd.ttf"), int(H * 0.017))
    fa = font(Path(r"C:\Windows\Fonts\tahoma.ttf"), int(H * 0.010))
    en_b = font(Path(r"C:\Windows\Fonts\calibrib.ttf"), int(H * 0.015))
    en = font(Path(r"C:\Windows\Fonts\calibri.ttf"), int(H * 0.012))

    tx = lx + lh + 8
    # Draw EN first on the right of FA visually: FA (navy) then AVID (cyan)
    name = reshape(COMPANY_FA)
    draw.text((tx, ly + 2), name, font=fa_b, fill=NAVY)
    nb = draw.textbbox((0, 0), name, font=fa_b)
    draw.text((tx + (nb[2] - nb[0]) + 6, ly + 4), COMPANY_EN, font=en_b, fill=CYAN)
    draw.text(
        (tx, ly + (nb[3] - nb[1]) + 6),
        reshape(TAGLINE),
        font=fa,
        fill=GREY,
    )

    img.paste(bars_r, (margin + stub_w + brand_w, y0), bars_r)

    # Sender / date
    y = y0 + deco_h + int(H * 0.025)
    draw.text((margin, y), reshape(SENDER_NAME), font=fa_b, fill=CYAN)
    y2 = y + int(H * 0.022)
    draw.text((margin, y2), reshape(SENDER_TITLE), font=fa, fill=GREY)
    meta_lines = [
        ("PHONE : ", PHONE),
        ("WEB : ", WEB),
        ("ADDR : ", ADDR_LINE1),
    ]
    yy = y2 + int(H * 0.02)
    for lab, val in meta_lines:
        draw.text((margin, yy), lab, font=en_b, fill=NAVY)
        lb = draw.textbbox((0, 0), lab, font=en_b)
        draw.text(
            (margin + lb[2] - lb[0] + 4, yy),
            reshape(val) if any("\u0600" <= ch <= "\u06FF" for ch in val) else val,
            font=fa,
            fill=SOFT,
        )
        yy += int(H * 0.018)

    date_txt = f"{DATE_LABEL} {DATE_VALUE}"
    db = draw.textbbox((0, 0), reshape(date_txt), font=fa)
    draw.text(
        (W - margin - (db[2] - db[0]), yy - int(H * 0.018)),
        reshape(date_txt),
        font=fa,
        fill=SOFT,
    )

    # Body greeting
    greet = reshape("با سلام و احترام،")
    draw.text(
        (W - margin, int(H * 0.28)),
        greet,
        font=fa,
        fill=NAVY,
        anchor="rt",
    )

    # Footer signatory + contacts
    footer_deco = Image.open(assets["footer_bars"]).convert("RGBA")
    fd_h = int(H * 0.075)
    fd_w = W - 2 * margin
    fd = footer_deco.resize((fd_w, fd_h), Image.Resampling.LANCZOS)
    fy = H - int(H * 0.04) - fd_h
    img.paste(fd, (margin, fy), fd)

    sy = fy - int(H * 0.12)
    draw.text((margin, sy), reshape(SENDER_NAME), font=fa_b, fill=CYAN)
    draw.text((margin, sy + int(H * 0.02)), reshape(SENDER_TITLE), font=fa, fill=GREY)
    draw.text(
        (margin, sy + int(H * 0.045)),
        "__________________",
        font=en,
        fill=NAVY,
    )

    icon_sz = int(H * 0.016)
    cx0 = int(W * 0.55)
    cy0 = sy
    contact_rows = [
        (assets["icon_phone"], f"{PHONE}  |  {MOBILE}"),
        (assets["icon_web"], f"{EMAIL}  |  {WEB}"),
        (assets["icon_location"], f"{ADDR_LINE1}  |  {ADDR_LINE2}"),
    ]
    for icon_p, text in contact_rows:
        ic = Image.open(icon_p).convert("RGBA").resize((icon_sz, icon_sz), Image.Resampling.LANCZOS)
        img.paste(ic, (cx0, cy0), ic)
        draw.text(
            (cx0 + icon_sz + 8, cy0),
            reshape(text) if any("\u0600" <= ch <= "\u06FF" for ch in text) else text,
            font=fa,
            fill=SOFT,
        )
        cy0 += icon_sz + 10

    out_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(out_path, "PNG", dpi=(150, 150))
    return out_path


def main() -> int:
    desktop_primary = Path(r"C:\Users\Almahdi Laptop\Desktop")
    desktop = desktop_primary if desktop_primary.is_dir() else resolve_desktop()

    print("NOTE=PSD not opened; geometry redrawn from JPG measurements")
    print("NAVY=#12374A CYAN=#03C2EF")

    print("Generating decorative PNGs…")
    assets = generate_decorative_assets()
    for k, p in assets.items():
        print(f"  {k}={p}")

    docx_path = desktop / "Letterhead-A4.docx"
    alt_docx = desktop / "Letterhead-A4-new.docx"
    preview_path = desktop / "Letterhead-A4-preview.png"

    print("Building editable DOCX…")
    try:
        build_docx(assets, docx_path)
        saved = docx_path
    except PermissionError:
        build_docx(assets, alt_docx)
        saved = alt_docx
        print("WARN=Target DOCX locked; wrote Letterhead-A4-new.docx")
        try:
            shutil.copy2(alt_docx, docx_path)
            saved = docx_path
            print("RETRY_OK=Overwrote Letterhead-A4.docx")
        except Exception:
            pass

    print("Building preview PNG…")
    try:
        build_preview_png(assets, preview_path)
        print(f"PREVIEW={preview_path}")
    except Exception as exc:
        print(f"PREVIEW_FAIL={exc}")

    print(f"SAVED={saved}")
    print(f"SIZE={saved.stat().st_size}")

    for desk in (
        Path(r"C:\Users\Almahdi Laptop\OneDrive\Desktop"),
        resolve_desktop(),
    ):
        if not desk.is_dir():
            continue
        try:
            if desk.resolve() == desktop.resolve():
                continue
        except Exception:
            continue
        for src, name in (
            (saved, "Letterhead-A4.docx"),
            (preview_path, "Letterhead-A4-preview.png"),
        ):
            if not src.is_file():
                continue
            alt = desk / name
            try:
                shutil.copy2(src, alt)
                print(f"COPIED={alt}")
            except PermissionError:
                print(f"SKIP_LOCKED={alt}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
